#!/usr/bin/env python3
"""Katte Bot — ingest admin documents into a searchable index.

Reads everything in ./docs (pdf, docx, xlsx, txt, md), chunks the text,
embeds each chunk with Ollama's nomic-embed-text, and writes index.json.
Re-run whenever documents change — that IS the "training" step.

Usage:  python3 ingest.py
Needs:  ollama running locally, `ollama pull nomic-embed-text` done once.
        pdf/docx/xlsx support: pip install pypdf python-docx openpyxl
        (txt/md work with no installs at all)
"""
import json, re, sys, urllib.request
from pathlib import Path

OLLAMA = "http://localhost:11434"
EMBED_MODEL = "nomic-embed-text"
DOCS = Path(__file__).parent / "docs"
INDEX = Path(__file__).parent / "index.json"
CHUNK_WORDS = 220
OVERLAP_WORDS = 50


def read_txt(p): return p.read_text(errors="ignore")

def read_pdf(p):
    try:
        from pypdf import PdfReader
    except ImportError:
        sys.exit(f"[skip-able] {p.name}: pip install pypdf for PDF support")
    return "\n".join(page.extract_text() or "" for page in PdfReader(p).pages)

def read_docx(p):
    try:
        import docx
    except ImportError:
        sys.exit(f"[skip-able] {p.name}: pip install python-docx for Word support")
    d = docx.Document(p)
    parts = [para.text for para in d.paragraphs]
    for t in d.tables:
        for row in t.rows:
            parts.append(" · ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)

def read_xlsx(p):
    try:
        import openpyxl
    except ImportError:
        sys.exit(f"[skip-able] {p.name}: pip install openpyxl for Excel support")
    wb = openpyxl.load_workbook(p, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[sheet: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells: parts.append(" · ".join(cells))
    return "\n".join(parts)

READERS = {".txt": read_txt, ".md": read_txt, ".pdf": read_pdf,
           ".docx": read_docx, ".xlsx": read_xlsx}


def chunk(text):
    words = re.sub(r"\s+", " ", text).split(" ")
    step = CHUNK_WORDS - OVERLAP_WORDS
    return [" ".join(words[i:i + CHUNK_WORDS])
            for i in range(0, max(len(words), 1), step)
            if words[i:i + CHUNK_WORDS]]


def embed(texts):
    req = urllib.request.Request(
        f"{OLLAMA}/api/embed",
        json.dumps({"model": EMBED_MODEL, "input": texts}).encode(),
        {"Content-Type": "application/json"})
    with urllib.request.urlopen(req, timeout=120) as r:
        return json.load(r)["embeddings"]


def main():
    files = sorted(p for p in DOCS.iterdir()
                   if p.suffix.lower() in READERS and not p.name.startswith("."))
    if not files:
        sys.exit(f"No documents found in {DOCS} — add pdf/docx/xlsx/txt/md files first.")
    entries = []
    for p in files:
        text = READERS[p.suffix.lower()](p).strip()
        if not text:
            print(f"  ! {p.name}: no extractable text (scanned pdf? needs OCR), skipping")
            continue
        title = text.splitlines()[0].lstrip("# ").strip()[:120] or p.stem
        chunks = chunk(text)
        # nomic-embed is trained with task prefixes; the title gives each
        # chunk page-level context so mid-page chunks stay findable
        vecs = embed([f"search_document: {title}\n{c}" for c in chunks])
        entries += [{"file": p.name, "title": title, "text": c, "vec": v}
                    for c, v in zip(chunks, vecs)]
        print(f"  ✓ {p.name}: {len(chunks)} chunk(s)")
    INDEX.write_text(json.dumps(entries))
    print(f"Indexed {len(entries)} chunks from {len(files)} file(s) → {INDEX.name}")


if __name__ == "__main__":
    main()
